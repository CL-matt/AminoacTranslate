import os
import re
import jieba
import json
import unicodedata
from pypinyin import pinyin, Style
import pdfplumber
import docx
from docx import Document

# Determine project root (AminoacTranslate)
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


class ToneStyle:
    def __init__(self, name, pypinyin_style, description=""):
        self.name = name
        self.pypinyin_style = pypinyin_style
        self.description = description


TONE_STYLES = [
    ToneStyle("古韵", Style.TONE, "保留声调符号，如 nǐ"),
    ToneStyle("音韵", Style.TONE2, "用数字表示声调，如 ni3"),
    ToneStyle("无声调", Style.NORMAL, "移除所有声调符号，如 ni")
]


def clean_pinyin(pinyin_str, style):
    if style.pypinyin_style == Style.NORMAL:
        cleaned = []
        for char in pinyin_str:
            if char == ' ':
                cleaned.append(char)
                continue
            normalized = unicodedata.normalize('NFD', char)
            for c in normalized:
                if not unicodedata.combining(c):
                    cleaned.append(c)
        return ''.join(cleaned).lower().replace("ü", "v")
    elif style.pypinyin_style == Style.TONE2:
        return re.sub(r'([a-zü]+)([1-4])', r'\1\2 ', pinyin_str.lower()).strip().replace("ü", "v")
    else:
        return pinyin_str.replace("ü", "v")


def reverse_pinyin_translation(chinese_text, tone_style, custom_dict=None):
    if custom_dict is None:
        custom_dict = {}
    if chinese_text in custom_dict:
        return custom_dict[chinese_text]

    words = jieba.cut(chinese_text, cut_all=False)
    pinyin_parts = []
    for word in words:
        if word in custom_dict:
            pinyin_parts.append(custom_dict[word])
        else:
            chars_pinyin = pinyin(word, style=tone_style.pypinyin_style)
            combined = ''.join([syllable[0] for syllable in chars_pinyin])
            pinyin_parts.append(combined)
    pinyin_str = ' '.join(pinyin_parts)
    processed = clean_pinyin(pinyin_str, tone_style)
    rev_parts = [seg[::-1] for seg in processed.split()[::-1]]
    out = ' '.join(rev_parts)
    return out[0].upper() + out[1:] if out else ""


def translate_with_punctuation(text: str, tone_style, custom_dict=None) -> str:
    if custom_dict is None:
        custom_dict = {}
    pattern = re.compile(r'([^。！？!?]*[。！？!?])')
    matches = pattern.findall(text)
    sentences = [match.strip() for match in matches if match.strip()]
    processed_sentences = []
    for sentence in sentences:
        content = sentence[:-1]
        delimiter = sentence[-1]
        words = list(jieba.cut(content))
        reversed_words = []
        for word in words:
            if word in custom_dict:
                py_str = custom_dict[word]
            else:
                pys = pinyin(word, style=tone_style.pypinyin_style)
                py_str = ''.join([p[0] for p in pys])
            cleaned = clean_pinyin(py_str, tone_style)
            reversed_word = cleaned[::-1]
            reversed_words.append(reversed_word)
        reversed_sentence = ''.join(
            word if re.match(r'[，。！？!?.,\s]', word) else f' {word}'
            for word in reversed_words[::-1]
        ).strip()
        processed_sentences.append(reversed_sentence + delimiter)
    processed_sentences = processed_sentences[::-1]
    result = ' '.join(processed_sentences).strip()
    if result:
        result_chars = list(result)
        capitalize_next = True
        for i, c in enumerate(result_chars):
            if c.isalpha() and capitalize_next:
                result_chars[i] = c.upper()
                capitalize_next = False
            elif c in ['。', '.', '！', '？']:
                capitalize_next = True
        result = ''.join(result_chars)
    return result


def translate_paragraphs(text: str, tone_style, custom_dict=None) -> str:
    paragraphs = text.split('\n')
    translated_paragraphs = []
    for paragraph in paragraphs:
        if paragraph.strip():
            translated_paragraph = translate_with_punctuation(paragraph.strip(), tone_style, custom_dict)
            translated_paragraphs.append(translated_paragraph)
    return '\n'.join(translated_paragraphs)


def detect_and_translate(text: str, tone_style, custom_dict=None) -> str:
    if '\n' in text.strip():
        return translate_paragraphs(text, tone_style, custom_dict)
    if re.search(r'[。！？!?.,]', text):
        return translate_with_punctuation(text, tone_style, custom_dict)
    return reverse_pinyin_translation(text, tone_style, custom_dict)


def translate_docx(input_path: str, output_path: str, tone_style, custom_dict=None):
    doc = Document(input_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if not para.text.strip():
                f.write('\n')
            else:
                translated = translate_with_punctuation(para.text, tone_style, custom_dict)
                f.write(translated + '\n')


def read_pdf(file_path):
    content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.extend(text.split('\n'))
    return [line.strip() for line in content if line.strip()]


def read_word(file_path):
    content = []
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text.strip())
    return content


def load_custom_dict(file_path=None):
    if file_path is None:
        file_path = os.path.join(PROJECT_ROOT, 'Aminoac', 'custom_dict.json')
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                custom_dict = json.load(f)
                for word in custom_dict.keys():
                    jieba.add_word(word)
                return custom_dict
        else:
            return {}
    except Exception:
        return {}


CUSTOM_TRANSLATION_DICT = load_custom_dict()


def split_long_text(text, max_length=50):
    sentences = re.split(r'(。|！|\!|\.|？|\?)', text)
    segments = []
    current_segment = ""
    for sentence in sentences:
        if len(current_segment) + len(sentence) <= max_length:
            current_segment += sentence
        else:
            segments.append(current_segment)
            current_segment = sentence
    if current_segment:
        segments.append(current_segment)
    if not segments:
        for i in range(0, len(text), max_length):
            segments.append(text[i:i + max_length])
    return segments
