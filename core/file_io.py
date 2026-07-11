"""
文件读写操作模块
"""
import json
import os
import re

import docx
import jieba
import pdfplumber
from docx import Document

from .translate import reverse_pinyin_translation, translate_with_punctuation

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _resolve_custom_dict_path(file_path=None):
    """根据项目结构解析自定义词典路径。"""
    if file_path:
        return file_path

    candidates = [
        os.path.join(PROJECT_ROOT, "Aminoac", "custom_dict.json"),
        os.path.join(PROJECT_ROOT, "custom_dict.json"),
        "custom_dict.json",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return os.path.join(PROJECT_ROOT, "Aminoac", "custom_dict.json")


def read_pdf(file_path):
    """读取 PDF 文件内容。"""
    content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.extend(text.split("\n"))
    return [line.strip() for line in content if line.strip()]


def read_word(file_path):
    """读取 Word 文件内容。"""
    content = []
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text.strip())
    return content


def load_custom_dict(file_path=None):
    """加载自定义翻译字典并更新 jieba 词典。"""
    resolved_path = _resolve_custom_dict_path(file_path)
    try:
        if os.path.exists(resolved_path):
            with open(resolved_path, "r", encoding="utf-8") as handle:
                custom_dict = json.load(handle)
                for word in custom_dict.keys():
                    jieba.add_word(word)
                return custom_dict
        return {}
    except Exception as exc:
        print(f"加载自定义字典失败: {exc}")
        return {}


def split_long_text(text, max_length=50):
    """将长文本按标点符号或固定长度分段。"""
    sentences = re.split(r"(。|！|\!|\.|？|\?)", text)
    segments = []
    current_segment = ""
    for sentence in sentences:
        if len(current_segment) + len(sentence) <= max_length:
            current_segment += sentence
        else:
            segments.append(current_segment)
            current_segment = sentence
    if current_segment:
        segments.append(current_segment)
    if not segments:
        for i in range(0, len(text), max_length):
            segments.append(text[i : i + max_length])
    return segments


def translate_docx(input_path: str, output_path: str, tone_style, custom_dict=None):
    """翻译 Word 文档。"""
    doc = Document(input_path)
    with open(output_path, "w", encoding="utf-8") as handle:
        for para in doc.paragraphs:
            if not para.text.strip():
                handle.write("\n")
            else:
                translated = translate_with_punctuation(para.text, tone_style, custom_dict)
                handle.write(translated + "\n")


def process_file(file_path: str, tone_style, custom_dict=None):
    """读取文件并返回翻译后的分段内容。"""
    if file_path.endswith(".pdf"):
        content = read_pdf(file_path)
    elif file_path.endswith(".docx"):
        content = read_word(file_path)
    else:
        try:
            import chardet

            with open(file_path, "rb") as handle:
                raw_data = handle.read()
                detected = chardet.detect(raw_data)
                encoding = detected["encoding"] if detected["encoding"] else "utf-8"
        except Exception:
            encoding = "utf-8"

        with open(file_path, "r", encoding=encoding, errors="replace") as handle:
            content = handle.read().split("\n")

    translated = []
    for para in content:
        if para.strip():
            segments = split_long_text(para)
            for segment in segments:
                translated.append(reverse_pinyin_translation(segment, tone_style, custom_dict))
    return translated